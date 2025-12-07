# app/utils/export_service.py

import base64
import io
from datetime import datetime
from typing import List, Tuple, Any, Optional

from sqlalchemy.orm import Session
from docx import Document  # pip install python-docx
from reportlab.pdfgen import canvas  # pip install reportlab

from app.models.json_document import JSONDocument
from app.models.export_template import ExportTemplate, ExportFormat
from app.models.field_config import FieldConfigSet, FieldConfig, ExportMaskType
from app.utils.mapping_engine import json_get


def _pick_template(
    db: Session,
    document: JSONDocument,
    export_format: ExportFormat,
    with_answers: bool,
    template_id: Optional[int],
) -> ExportTemplate:
    if template_id:
        t = db.query(ExportTemplate).get(template_id)
        if not t:
            raise ValueError("Template not found")
        return t

    t = (
        db.query(ExportTemplate)
        .filter(
            ExportTemplate.json_type_id == document.json_type_id,
            ExportTemplate.format == export_format,
            ExportTemplate.with_answers == with_answers,
            ExportTemplate.is_active == True,
        )
        .order_by(ExportTemplate.created_at.desc())
        .first()
    )
    if not t:
        raise ValueError("No matching export template found")
    return t


def _pick_field_config_set(db: Session, json_type_id: int) -> Optional[FieldConfigSet]:
    default_set = (
        db.query(FieldConfigSet)
        .filter(
            FieldConfigSet.json_type_id == json_type_id,
            FieldConfigSet.is_default == True,
        )
        .order_by(FieldConfigSet.created_at.desc())
        .first()
    )
    if default_set:
        return default_set

    return (
        db.query(FieldConfigSet)
        .filter(FieldConfigSet.json_type_id == json_type_id)
        .order_by(FieldConfigSet.created_at.asc())
        .first()
    )


def _build_field_values(
    db: Session,
    document: JSONDocument,
    with_answers: bool,
) -> List[Tuple[str, str]]:
    """
    Returns list of (label, value_str) according to field_config.
    If with_answers is False, values are blanked.
    """
    cfg_set = _pick_field_config_set(db, document.json_type_id)
    if not cfg_set:
        # Fallback: single block dumping the JSON.
        return [("JSON", _safe_stringify(document.raw_json if with_answers else {}))]

    configs: List[FieldConfig] = (
        db.query(FieldConfig)
        .filter(
            FieldConfig.config_set_id == cfg_set.id,
            FieldConfig.show_in_export == True,
        )
        .order_by(FieldConfig.order_index.asc())
        .all()
    )

    rows: List[Tuple[str, str]] = []
    for cfg in configs:
        label = cfg.label or cfg.json_path

        if not with_answers:
            val_str = ""
        else:
            raw_val = json_get(document.raw_json, cfg.json_path)
            val_str = _apply_mask_and_stringify(raw_val, cfg.export_mask_type)

        rows.append((label, val_str))

    return rows


def _apply_mask_and_stringify(value: Any, mask_type: ExportMaskType) -> str:
    if mask_type == ExportMaskType.HIDE_VALUE:
        return ""
    if mask_type == ExportMaskType.REDACT:
        return "***"
    return _safe_stringify(value)


def _safe_stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (str, int, float, bool)):
        return str(value)
    try:
        return json_dumps_compact(value)
    except Exception:
        return str(value)


def json_dumps_compact(value: Any) -> str:
    import json

    return json.dumps(value, ensure_ascii=False, separators=(",", ":"))


def _create_docx(
    fields: List[Tuple[str, str]],
    doc_title: str,
) -> bytes:
    doc = Document()
    if doc_title:
        doc.add_heading(doc_title, level=1)

    for label, value in fields:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        if value:
            p.add_run(str(value))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _create_pdf(
    fields: List[Tuple[str, str]],
    doc_title: str,
) -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer)

    y = 800

    if doc_title:
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, y, doc_title)
        y -= 30

    c.setFont("Helvetica", 10)

    for label, value in fields:
        line = f"{label}: {value}" if value else f"{label}:"
        if y < 50:
            c.showPage()
            y = 800
            c.setFont("Helvetica", 10)
        c.drawString(50, y, line)
        y -= 20

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer.getvalue()


def generate_export(
    db: Session,
    document: JSONDocument,
    export_format: ExportFormat,
    with_answers: bool,
    template_id: Optional[int] = None,
) -> tuple[str, str]:
    """
    Returns (file_data_base64, file_name)
    """
    template = _pick_template(
        db=db,
        document=document,
        export_format=export_format,
        with_answers=with_answers,
        template_id=template_id,
    )

    fields = _build_field_values(db, document, with_answers)

    title = template.name or "Export"
    timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")

    if export_format == ExportFormat.DOCX:
        raw_bytes = _create_docx(fields, title)
        file_name = f"{title}_{timestamp}.docx"
    elif export_format == ExportFormat.PDF:
        raw_bytes = _create_pdf(fields, title)
        file_name = f"{title}_{timestamp}.pdf"
    else:
        raise ValueError("Unsupported export format")

    encoded = base64.b64encode(raw_bytes).decode("utf-8")
    return encoded, file_name
